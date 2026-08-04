from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
XLSX_PATH = ROOT / "storage" / "templates" / "Lembar Kerja 095-152.xlsx"
DOCX_PATH = (
    ROOT
    / "storage"
    / "templates"
    / "docx"
    / "CCI-KAL-FOM-152 LEMBAR KERJA KALIBRASI TORQUE GAUGE (Rev 00).docx"
)
OUTPUT_PATH = ROOT / "storage" / "generated" / "torque-template-comparison.json"


def normalized(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


document = Document(DOCX_PATH)


def extract_tables(tables: object, collected_text: list[str]) -> list[list[list[str]]]:
    extracted: list[list[list[str]]] = []
    for table in tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [normalized(cell.text) for cell in row.cells]
            rows.append(cells)
            collected_text.extend(text for text in cells if text)
        extracted.append(rows)
    return extracted


docx_paragraphs = [normalized(paragraph.text) for paragraph in document.paragraphs]
docx_paragraphs = [text for text in docx_paragraphs if text]
docx_text: list[str] = list(docx_paragraphs)
docx_tables = extract_tables(document.tables, docx_text)

headers: list[dict[str, object]] = []
footers: list[dict[str, object]] = []
for section_number, section in enumerate(document.sections, start=1):
    for destination, part in [(headers, section.header), (footers, section.footer)]:
        part_paragraphs = [normalized(paragraph.text) for paragraph in part.paragraphs]
        part_paragraphs = [text for text in part_paragraphs if text]
        docx_text.extend(part_paragraphs)
        part_tables = extract_tables(part.tables, docx_text)
        destination.append(
            {
                "section": section_number,
                "paragraphs": part_paragraphs,
                "tables": part_tables,
            }
        )

workbook = load_workbook(XLSX_PATH, read_only=False, data_only=False)
worksheet = workbook["Torque Gauge"]
excel_cells = {
    cell.coordinate: normalized(cell.value)
    for row in worksheet.iter_rows()
    for cell in row
    if cell.value is not None
}
excel_text = list(excel_cells.values())

code_pattern = re.compile(r"CCI-KAL-(?:FOM|WI)-[0-9A-Z]+", re.IGNORECASE)
revision_pattern = re.compile(r"(?:REV(?:ISI)?|REV\.)\s*[:\-]?\s*0*([0-9]+)", re.IGNORECASE)


def codes(values: list[str]) -> list[str]:
    return sorted({match.group(0).upper() for value in values for match in code_pattern.finditer(value)})


def revisions(values: list[str]) -> list[str]:
    found = set()
    for value in values:
        for match in revision_pattern.finditer(value):
            found.add(match.group(1))
    return sorted(found)


docx_codes = codes(docx_text)
excel_codes = codes(excel_text)
docx_revisions = revisions(docx_text)
excel_revisions = revisions(excel_text)

comparison = {
    "sources": {"xlsx": str(XLSX_PATH), "docx": str(DOCX_PATH)},
    "excel": {
        "sheet": worksheet.title,
        "dimensions": worksheet.calculate_dimension(),
        "nonEmptyCellCount": len(excel_cells),
        "codes": excel_codes,
        "revisionsDetected": excel_revisions,
        "notableCells": {
            coordinate: excel_cells.get(coordinate)
            for coordinate in ["C2", "C3", "C4", "E2", "F9", "F46"]
        },
    },
    "docx": {
        "paragraphCount": len(document.paragraphs),
        "tableCount": len(document.tables),
        "nonEmptyTextCount": len(docx_text),
        "codes": docx_codes,
        "revisionsDetected": docx_revisions,
        "paragraphs": docx_paragraphs,
        "tables": docx_tables,
        "headers": headers,
        "footers": footers,
    },
    "comparison": {
        "codesOnlyInExcel": sorted(set(excel_codes) - set(docx_codes)),
        "codesOnlyInDocx": sorted(set(docx_codes) - set(excel_codes)),
        "revisionSetsMatch": excel_revisions == docx_revisions,
        "excelTitleMentionsTorque": any("TORQUE" in text.upper() for text in excel_text),
        "docxTitleMentionsTorque": any("TORQUE" in text.upper() for text in docx_text),
        "excelTitleMentionsDigitalPressure": any(
            "DIGITAL PRESSURE" in text.upper() for text in excel_text
        ),
        "docxTitleMentionsDigitalPressure": any(
            "DIGITAL PRESSURE" in text.upper() for text in docx_text
        ),
    },
}

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH.write_text(json.dumps(comparison, ensure_ascii=False, indent=2), encoding="utf-8")

print(json.dumps({"output": str(OUTPUT_PATH), **comparison["comparison"], "excel": comparison["excel"], "docxSummary": {key: comparison["docx"][key] for key in ["paragraphCount", "tableCount", "nonEmptyTextCount", "codes", "revisionsDetected"]}}, ensure_ascii=False, indent=2))
